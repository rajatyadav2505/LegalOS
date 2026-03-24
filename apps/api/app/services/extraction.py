from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import pytesseract
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from PIL import Image
from pypdf import PdfReader

from app.core.config import get_settings


@dataclass(slots=True)
class ExtractedParagraph:
    text: str
    paragraph_number: int
    page_number: int | None = None
    heading: str | None = None


@dataclass(slots=True)
class ExtractedDocument:
    title: str
    paragraphs: list[ExtractedParagraph]
    extraction_method: str

    @property
    def full_text(self) -> str:
        return "\n\n".join(paragraph.text for paragraph in self.paragraphs)


class DocumentExtractor:
    def __init__(self) -> None:
        self.settings = get_settings()

    def extract(self, *, file_name: str, content_type: str, payload: bytes) -> ExtractedDocument:
        suffix = Path(file_name).suffix.lower()
        if suffix in {".txt", ".md", ".rtf"} or content_type.startswith("text/"):
            return self._from_text(file_name, self._decode_text(payload), extraction_method="text")
        if suffix in {".html", ".htm"} or "html" in content_type:
            soup = BeautifulSoup(self._decode_text(payload), "html.parser")
            return self._from_text(file_name, soup.get_text("\n"), extraction_method="html_text")
        if suffix == ".docx" or "wordprocessingml" in content_type:
            doc = DocxDocument(BytesIO(payload))
            text = "\n\n".join(
                paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()
            )
            return self._from_text(file_name, text, extraction_method="docx")
        if suffix == ".pdf" or content_type == "application/pdf":
            reader = PdfReader(BytesIO(payload))
            paragraphs: list[ExtractedParagraph] = []
            counter = 1
            used_ocr = False
            for page_number, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                page_paragraphs = self._split_paragraphs(page_text)
                if not page_paragraphs:
                    ocr_text = self._ocr_pdf_page(page)
                    if ocr_text:
                        page_paragraphs = self._split_paragraphs(ocr_text)
                        used_ocr = True
                for paragraph in page_paragraphs:
                    paragraphs.append(
                        ExtractedParagraph(
                            text=paragraph,
                            paragraph_number=counter,
                            page_number=page_number,
                        )
                    )
                    counter += 1
            return ExtractedDocument(
                title=Path(file_name).stem,
                paragraphs=paragraphs,
                extraction_method="pdf_text_ocr" if used_ocr else "pdf_text",
            )
        if (
            suffix in {".png", ".jpg", ".jpeg", ".tiff", ".tif"}
            or content_type.startswith("image/")
        ):
            image = Image.open(BytesIO(payload))
            try:
                text = pytesseract.image_to_string(
                    image,
                    lang=self.settings.ocr_languages,
                    config=self.settings.ocr_tesseract_config,
                )
            except pytesseract.TesseractNotFoundError as exc:
                raise ValueError("OCR is unavailable because Tesseract is not installed") from exc
            return self._from_text(file_name, text, extraction_method="image_ocr")
        raise ValueError(f"Unsupported file type for extraction: {file_name}")

    def _from_text(
        self,
        file_name: str,
        text: str,
        *,
        extraction_method: str,
    ) -> ExtractedDocument:
        paragraphs = [
            ExtractedParagraph(text=paragraph, paragraph_number=index)
            for index, paragraph in enumerate(self._split_paragraphs(text), start=1)
        ]
        return ExtractedDocument(
            title=Path(file_name).stem,
            paragraphs=paragraphs,
            extraction_method=extraction_method,
        )

    @staticmethod
    def _decode_text(payload: bytes) -> str:
        for encoding in ("utf-8", "utf-16", "latin-1"):
            try:
                return payload.decode(encoding)
            except UnicodeDecodeError:
                continue
        return payload.decode("utf-8", errors="ignore")

    def _ocr_pdf_page(self, page) -> str:
        extracted: list[str] = []
        for image_file in getattr(page, "images", []):
            image_data = getattr(image_file, "data", None)
            if not image_data:
                continue
            image = Image.open(BytesIO(image_data))
            try:
                text = pytesseract.image_to_string(
                    image,
                    lang=self.settings.ocr_languages,
                    config=self.settings.ocr_tesseract_config,
                )
            except pytesseract.TesseractNotFoundError as exc:
                raise ValueError("OCR is unavailable because Tesseract is not installed") from exc
            if text.strip():
                extracted.append(text)
        return "\n\n".join(extracted)

    @staticmethod
    def _split_paragraphs(text: str) -> list[str]:
        paragraphs = [block.strip() for block in text.replace("\r\n", "\n").split("\n\n")]
        return [paragraph for paragraph in paragraphs if paragraph]
